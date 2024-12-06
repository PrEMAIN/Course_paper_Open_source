import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    # Выравниваем параграф по центру
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Добавляем динамическую нумерацию страниц
    page_num_run = paragraph.add_run()
    
    # Создаем XML-элементы для поля нумерации
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"  # Устанавливаем для поля текст "PAGE" (нумерация страницы)
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    
    # Добавляем все части поля в run
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    
    # Устанавливаем шрифт Times New Roman для номера страницы
    page_num_run.font.name = 'Times New Roman'
    page_num_run.font.size = Pt(14)

def process_document(input_path, output_path):
    # Открываем документ
    doc = Document(input_path)

    # Устанавливаем минимальные поля документа
    section = doc.sections[0]
    section.top_margin = Inches(0.8)   # Верхнее поле — 20 мм
    section.bottom_margin = Inches(0.8)  # Нижнее поле — 20 мм
    section.left_margin = Inches(0.8)   # Левое поле — 20 мм
    section.right_margin = Inches(0.4)  # Правое поле — 10 мм

    # Применяем жирный шрифт и черный цвет для всех стилей заголовков (Heading)
    for style in doc.styles:
        if style.name.startswith('Heading'):  # Проверяем, что стиль заголовка
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)  # Размер шрифта для заголовков
            style.font.bold = True  # Жирный шрифт
            style.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет текста

    # Применяем изменения ко всем параграфам
    for paragraph in doc.paragraphs:
        # Выравнивание текста
        if paragraph.style.name.startswith('Heading'):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Заголовки по левому краю
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Текст по ширине

        # Абзацный отступ
        paragraph.paragraph_format.first_line_indent = Inches(0.5)  # 1.25 см
        
        # Межстрочный интервал
        paragraph.paragraph_format.line_spacing = 1.5
        
        # Применение шрифта и размера для обычного текста
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    # Настроим колонтитулы для первой страницы и всех остальных
    section.different_first_page = False  # Используем одинаковые колонтитулы для всех страниц
    section.starting_number = 2  # Начинаем нумерацию с 2

    # Колонтитул для всех страниц
    footer_other_pages = section.footer
    paragraph = footer_other_pages.paragraphs[0]
    add_page_number(paragraph)  # Добавляем нумерацию для всех страниц

    # Сохраняем изменения
    doc.save(output_path)

def open_file():
    # Открываем диалог выбора файла
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        try:
            # Указываем путь для сохранения результата
            output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
            if output_path:
                process_document(file_path, output_path)
                messagebox.showinfo("Успех", f"Документ сохранен: {output_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")
    else:
        messagebox.showwarning("Предупреждение", "Файл не выбран!")

def create_app():
    # Создаем основное окно приложения
    root = tk.Tk()
    root.title("Нормоконтроль документов")
    
    # Устанавливаем размер окна
    root.geometry("400x200")

    # Создаем кнопку для загрузки и обработки файла
    button = tk.Button(root, text="Загрузить документ", command=open_file)
    button.pack(pady=50)

    # Запуск приложения
    root.mainloop()

if __name__ == "__main__":
    create_app()
